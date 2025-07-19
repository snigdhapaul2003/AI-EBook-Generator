import json
from datetime import datetime
from typing import Dict, List, Any, Tuple, Optional
from google import genai
from google.genai import types
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fpdf import FPDF
import re
import time

from config import (
    EbookState, Chapter, ChapterStatus, OutlineStatus,
    FormatType, PROMPTS, GenerationConfig
)

class GeminiGenerator:
    """Helper class for generating content with Gemini."""

    def __init__(self, api_key: str, config: GenerationConfig):
        """Initialize the Gemini client."""
        self.api_key = api_key
        self.config = config
        self.client = genai.Client(api_key=self.api_key)

    def generate_content(self, prompt: str, stream: bool = False) -> str:
        """Generate content using the Gemini API with proper syntax."""
        model = self.config.model_name

        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(text=prompt),
                ],
            ),
        ]

        generate_content_config = types.GenerateContentConfig(
            temperature=self.config.temperature,
            top_p=self.config.top_p,
            top_k=self.config.top_k,
            max_output_tokens=self.config.max_output_tokens,
            thinking_config=types.ThinkingConfig(
                thinking_budget=0,
            ),
            response_mime_type="text/plain",
        )

        if stream:
            response_text = ""
            print("Generating content...")
            for chunk in self.client.models.generate_content_stream(
                model=model,
                contents=contents,
                config=generate_content_config,
            ):
                if hasattr(chunk, 'text'):
                    if chunk.text:  # Make sure it's not None or empty
                        response_text += chunk.text
            print("\n")
            return response_text
        else:
            response = self.client.models.generate_content(
                model=model,
                contents=contents,
                config=generate_content_config,
            )
            return response.text


class InitializeNode:
    """Node to initialize the e-book generation process"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Initialize the e-book state with user input"""
        topic = state.get("topic", "")
        target_audience = state.get("target_audience", "general readers")
        tone = state.get("tone", "professional but conversational")
        format_type = state.get("format_type", FormatType.DOC)
        additional_description = state.get("additional_description", "")

        # Create initial state
        ebook_state = EbookState(
            topic=topic,
            target_audience=target_audience,
            tone=tone,
            format_type=format_type,
            additional_description=additional_description,
            creation_date=datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        )

        # Update the graph state
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Initializing node")
        return state


class OutlineGenerationNode:
    """Node to generate the e-book outline"""

    def __init__(self, gemini: GeminiGenerator):
        self.gemini = gemini

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Generate the outline for the e-book"""
        ebook_state = EbookState(**state["ebook_state"])

        # Prepare additional requirements text
        additional_requirements = ""
        if ebook_state.additional_description and ebook_state.additional_description.strip():
            additional_requirements = f"\nAdditional requirements: {ebook_state.additional_description}"

        # Create prompt for outline generation
        prompt = PROMPTS["outline_generation"].format(
            topic=ebook_state.topic,
            target_audience=ebook_state.target_audience,
            tone=ebook_state.tone,
            additional_description=additional_requirements
        )

        # Generate outline using Gemini
        outline_json = self.gemini.generate_content(prompt)

        try:
            # Extract JSON content
            if "```json" in outline_json:
                outline_json = outline_json.split("```json")[1].split("```")[0].strip()
            elif "```" in outline_json:
                outline_json = outline_json.split("```")[1].split("```")[0].strip()

            outline_data = json.loads(outline_json)

            # Update ebook state with outline data
            ebook_state.outline = outline_data
            ebook_state.title = outline_data["title"]
            ebook_state.chapters = [Chapter(**chapter) for chapter in outline_data["chapters"]]
            ebook_state.outline_status = OutlineStatus.REVIEW

            # Update the graph state
            state["ebook_state"] = ebook_state.model_dump()
            print("I am in Outline Generation node")
            return state

        except Exception as e:
            # Handle JSON parsing error
            state["error"] = f"Error parsing outline: {str(e)}"
            return state


class OutlineReviewNode:
    """Node to automatically review the outline and decide if revision is needed."""

    def __init__(self, gemini):
        self.gemini = gemini  # Pass your LLM object (Gemini or other)

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Evaluate outline and either approve or mark for revision based on scores."""
        ebook_state = EbookState(**state["ebook_state"])
        outline = ebook_state.outline

        prompt = PROMPTS["outline_review"].format(outline=outline)

        # === LLM CALL ===
        try:
            response = self.gemini.generate_content(prompt)
            review_text = response.strip()
        except Exception as e:
            print(f"LLM review failed: {e}")
            review_text = ""

        # Extract scores
        import re
        scores = re.findall(r":\s*(\d{1,2})/10", review_text)
        scores = [int(s) for s in scores if s.isdigit()]
        total_score = sum(scores)
        average_score = total_score / len(scores) if scores else 0

        # Threshold to decide if it's good enough
        passing_threshold = 50  # out of 70

        print(f"Total Score: {total_score}, Average Score: {average_score}")

        if total_score >= passing_threshold:
            ebook_state.outline_status = OutlineStatus.APPROVED
        else:
            ebook_state.outline_status = OutlineStatus.REVISING
            ebook_state.revision_notes = review_text

        # Update state
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Outline review Node")
        return state


class OutlineRevisionNode:
    """Node to revise the e-book outline based on feedback"""

    def __init__(self, gemini: GeminiGenerator):
        self.gemini = gemini

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Revise the outline based on user feedback"""
        ebook_state = EbookState(**state["ebook_state"])

        # Create prompt for outline revision
        revision_prompt = f"""
        Revise the outline for the e-book "{ebook_state.title}" on the topic "{ebook_state.topic}"
        based on the following feedback:

        {ebook_state.revision_notes}

        Current outline:
        {json.dumps({"title": ebook_state.title, "chapters": [chapter.model_dump() for chapter in ebook_state.chapters]}, indent=2)}

        Format the output as JSON with the following structure:
        {{
            "title": "E-book Title",
            "chapters": [
                {{
                    "chapter_number": 1,
                    "title": "Chapter Title",
                    "bullet_points": ["Point 1", "Point 2", "Point 3"]
                }}
            ]
        }}
        """

        # Generate revised outline
        revised_outline_json = self.gemini.generate_content(revision_prompt)

        try:
            # Extract JSON content
            if "```json" in revised_outline_json:
                revised_outline_json = revised_outline_json.split("```json")[1].split("```")[0].strip()
            elif "```" in revised_outline_json:
                revised_outline_json = revised_outline_json.split("```")[1].split("```")[0].strip()

            revised_outline_data = json.loads(revised_outline_json)

            # Update ebook state with revised outline data
            ebook_state.title = revised_outline_data["title"]
            ebook_state.chapters = [Chapter(**chapter) for chapter in revised_outline_data["chapters"]]
            ebook_state.outline_status = OutlineStatus.REVIEW
            ebook_state.revision_notes = None

            # Update the graph state
            state["ebook_state"] = ebook_state.model_dump()
            state.pop("revision_requested", None)
            print("I am in Outline Revision node")
            return state

        except Exception as e:
            # Handle JSON parsing error
            state["error"] = f"Error parsing revised outline: {str(e)}"
            return state


class ContextManagerNode:
    """Node to prepare context for chapter generation"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context from previous chapters for the current chapter"""
        ebook_state = EbookState(**state["ebook_state"])

        current_idx = ebook_state.current_chapter_index

        # Update the current chapter status
        if current_idx < len(ebook_state.chapters):
            ebook_state.chapters[current_idx].status = ChapterStatus.GENERATING

        # Build context from previous chapters
        previous_context = ""
        if current_idx > 0:
            previous_context = "Previous chapters covered:\n"
            for idx in range(current_idx):
                chapter = ebook_state.chapters[idx]
                if chapter.content:
                    # Create a summary of previous content
                    previous_context += f"- Chapter {chapter.chapter_number}: {chapter.title} (Summary of key points)\n"
                    # We could use an LLM to generate a better summary here
                else:
                    previous_context += f"- Chapter {chapter.chapter_number}: {chapter.title}\n"
                    for point in chapter.bullet_points:
                        previous_context += f"  • {point}\n"

        # Update the state
        state["ebook_state"] = ebook_state.model_dump()
        state["previous_context"] = previous_context
        print("I am in Context Manager node")
        return state


class ChapterGenerationNode:
    """Node to generate a chapter"""

    def __init__(self, gemini: GeminiGenerator):
        self.gemini = gemini

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for the current chapter"""
        ebook_state = EbookState(**state["ebook_state"])
        previous_context = state.get("previous_context", "")

        current_idx = ebook_state.current_chapter_index

        if current_idx >= len(ebook_state.chapters):
            state["error"] = f"Chapter index {current_idx} out of range"
            return state

        current_chapter = ebook_state.chapters[current_idx]

        # Progress callback with chapter info
        try:
            import streamlit as st
            if hasattr(st.session_state, 'progress_callback'):
                chapter_info = f"Chapter {current_idx + 1}: {current_chapter.title}"
                st.session_state.progress_callback("generate_chapter", chapter_info)
        except:
            pass

        # Prepare bullet points for prompt
        bullet_points = "\n".join([f"- {point}" for point in current_chapter.bullet_points])

        # Prepare additional context if available
        additional_context = ""
        if ebook_state.additional_description and ebook_state.additional_description.strip():
            additional_context = f"\n\nAdditional context to consider: {ebook_state.additional_description}"

        # Create prompt for chapter generation
        prompt = PROMPTS["chapter_generation"].format(
            chapter_number=current_chapter.chapter_number,
            chapter_title=current_chapter.title,
            ebook_title=ebook_state.title,
            previous_context=previous_context,
            bullet_points=bullet_points,
            additional_context=additional_context,
            target_audience=ebook_state.target_audience,
            tone=ebook_state.tone
        )

        # Generate chapter content with streaming
        chapter_content = self.gemini.generate_content(prompt)

        # Remove all asterisk characters
        chapter_content = "".join(chapter_content).replace("*", "")

        # print(chapter_content)

        # Update chapter status and content
        ebook_state.chapters[current_idx].content = chapter_content
        ebook_state.chapters[current_idx].status = ChapterStatus.REVIEW

        # Update the graph state
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Chapter Generation node")
        return state

class ChapterReviewNode:
    """Node to handle user review of a chapter"""

    def __init__(self, gemini):
        self.gemini = gemini  # Pass your LLM object (Gemini or other)

    def extract_json_block(self, text: str) -> str:
        """
        Extracts the first valid JSON block from the text, even if it's wrapped in triple backticks.
        """
        # Remove triple backtick code blocks, e.g. ```json ... ```
        json_match = re.search(r"```(?:json)?(.*?)```", text, re.DOTALL)
        if json_match:
            cleaned = json_match.group(1).strip()
        else:
            # Fallback: Try to find a raw JSON object by matching braces
            brace_match = re.search(r"(\{.*\})", text, re.DOTALL)
            cleaned = brace_match.group(1).strip() if brace_match else text.strip()

        return cleaned

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Handle user review of the current chapter"""
        ebook_state = EbookState(**state["ebook_state"])
        current_idx = ebook_state.current_chapter_index

        previous_context = state.get("previous_context", "")
        current_chapter = ebook_state.chapters[current_idx]
        current_chapter_content = current_chapter.content
        bullet_points = "\n".join([f"- {point}" for point in current_chapter.bullet_points])

        prompt = PROMPTS["chapter_review"].format(previous_context=previous_context, bullet_points=bullet_points, chapter_text=current_chapter_content)

        try:
            time.sleep(5)
            raw_response = self.gemini.generate_content(prompt)
            cleaned_json = self.extract_json_block(raw_response)
            review_data = json.loads(cleaned_json)
            # print("Parsed LLM Review JSON:", review_data)
        except Exception as e:
            print(f"LLM review failed or malformed JSON: {e}")
            review_data = {
                "needs_revision": False,
                "quality_score": 0.0,
                "tone": "unknown",
                "issues": [],
                "revision_suggestions": []
            }


        if review_data.get("needs_revision", False):
            print("Chapter revision requested based on LLM review")
            current_chapter.status = ChapterStatus.REVISING
            current_chapter.revision_notes = "\n".join(review_data.get("revision_suggestions", []))
            current_chapter.revision_count += 1
        elif review_data.get("quality_score")<8.0:
            print("Chapter revision requested based on LLM review")
            current_chapter.status = ChapterStatus.REVISING
            current_chapter.revision_notes = "\n".join(review_data.get("revision_suggestions", []))
            current_chapter.revision_count += 1
        else:
            current_chapter.status = ChapterStatus.COMPLETED

        # Update the graph state
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Chapter Review node")
        return state


class ChapterRevisionNode:
    """Node to revise a chapter based on feedback"""

    def __init__(self, gemini: GeminiGenerator):
        self.gemini = gemini

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Revise the current chapter based on user feedback"""
        ebook_state = EbookState(**state["ebook_state"])
        current_idx = ebook_state.current_chapter_index

        if current_idx >= len(ebook_state.chapters):
            state["error"] = f"Chapter index {current_idx} out of range"
            return state

        current_chapter = ebook_state.chapters[current_idx]

        # Create prompt for chapter revision
        revision_prompt = PROMPTS["chapter_revision"].format(chapter_number=current_chapter.chapter_number, chapter_title=current_chapter.title, ebook_title=ebook_state.title, revision_notes=current_chapter.revision_notes, content=current_chapter.content)

        # Generate revised chapter content
        revised_content = self.gemini.generate_content(revision_prompt)

        revised_content = "".join(revised_content).replace("*", "")

        # Update chapter
        ebook_state.chapters[current_idx].content = revised_content
        ebook_state.chapters[current_idx].status = ChapterStatus.REVIEW
        ebook_state.chapters[current_idx].revision_notes = None

        # Update the graph state
        state["ebook_state"] = ebook_state.model_dump()
        state.pop("chapter_revision_requested", None)
        print("I am in Chapter Revision node")
        return state


class ChapterCompletionNode:
    """Node to check if all chapters are complete"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Check if all chapters are complete or move to next chapter"""
        ebook_state = EbookState(**state["ebook_state"])
        current_idx = ebook_state.current_chapter_index

        # Mark current chapter as complete if it's not already
        if current_idx < len(ebook_state.chapters):
            if ebook_state.chapters[current_idx].status != ChapterStatus.COMPLETED:
                ebook_state.chapters[current_idx].status = ChapterStatus.COMPLETED

        # IMPORTANT: Check if we have more chapters to generate
        if current_idx + 1 < len(ebook_state.chapters):
            # Move to next chapter
            ebook_state.current_chapter_index += 1
            state["has_more_chapters"] = True
            print(f"Moving to next chapter: {ebook_state.current_chapter_index+1}")
        else:
            # All chapters complete
            state["has_more_chapters"] = False
            print("All chapters complete, proceeding to compilation")

        # Update the graph state
        state["ebook_state"] = ebook_state.model_dump()  # Use model_dump instead of dict
        print("I am in Chapter Completion node")
        return state


class CompilationNode:
    """Node to compile all chapters into a complete e-book"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Compile all chapters into a complete document"""
        ebook_state = EbookState(**state["ebook_state"])

        # Progress callback
        try:
            import streamlit as st
            if hasattr(st.session_state, 'progress_callback'):
                st.session_state.progress_callback("compilation")
        except:
            pass

        # Combine all chapters into a single document
        compiled_content = f"# {ebook_state.title}\n\n"

        # Add metadata
        compiled_content += f"*Created: {ebook_state.creation_date}*\n\n"

        # Add table of contents
        compiled_content += "## Table of Contents\n\n"
        for chapter in ebook_state.chapters:
            compiled_content += f"- Chapter {chapter.chapter_number}: {chapter.title}\n"
        compiled_content += "\n\n"

        # Add chapters
        for chapter in ebook_state.chapters:
            compiled_content += f"## Chapter {chapter.chapter_number}: {chapter.title}\n\n"
            if chapter.content:
                compiled_content += chapter.content + "\n\n"
            else:
                compiled_content += "*Content not generated*\n\n"

        # Add the compiled content to the state
        state["compiled_content"] = compiled_content
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Compilation node")
        return state


class FormatConversionNode:
    """Node to convert the compiled e-book to the requested format"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Convert the compiled e-book to the requested format"""

        ebook_state = EbookState(**state["ebook_state"])
        compiled_content = state.get("compiled_content", "")
        output_filename = ""

        # Get the desired format
        format_type = ebook_state.format_type
        print(f"Converting to format: {format_type}")
        title = ebook_state.title
        
        # Sanitize filename to remove invalid characters
        def sanitize_filename(filename):
            """Remove or replace invalid characters for Windows filenames"""
            # List of invalid characters for Windows filenames
            invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
            
            # Replace invalid characters with underscores
            for char in invalid_chars:
                filename = filename.replace(char, '_')
            
            # Replace multiple spaces/underscores with single underscore
            import re
            filename = re.sub(r'[_\s]+', '_', filename)
            
            # Remove leading/trailing underscores and spaces
            filename = filename.strip('_ ')
            
            # Ensure filename isn't empty
            if not filename:
                filename = "untitled"
            
            # Limit length to avoid path too long errors
            if len(filename) > 100:
                filename = filename[:100]
            
            return filename
        
        filename_base = sanitize_filename(title)
        print(f"Original title: {title}")
        print(f"Sanitized filename base: {filename_base}")
        
        if filename_base != title.replace(" ", "_"):
            print("⚠️  Note: Filename was sanitized to remove invalid characters (like : | ? * < > \" / \\)")
        
        output_dir = "./output/"

        class PDF(FPDF):
            pass

        def clean_text(text):
            # Optional: Replace fancy quotes and dashes with basic ones
            return (
                text.replace("’", "'")
                    .replace("“", '"')
                    .replace("”", '"')
                    .replace("–", "-")
                    .replace("—", "--")
                    .encode("latin-1", errors="ignore")  # Strip non-latin-1 characters
                    .decode("latin-1")
            )

        if format_type == FormatType.MARKDOWN:
            print("Saving as Markdown")
            output_filename = f"{filename_base}.md"
            
            # Ensure output directory exists
            os.makedirs("./output", exist_ok=True)
            
            # Use absolute path for better compatibility
            output_path = os.path.abspath(f"./output/{output_filename}")
            print(f"Saving Markdown file to: {output_path}")
            
            try:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(f"# {title}\n\n")
                    f.write("## Table of Contents\n\n")
                    for chapter in ebook_state.chapters:
                        f.write(f"- Chapter {chapter.chapter_number}: {chapter.title}\n")
                    f.write("\n")

                    for chapter in ebook_state.chapters:
                        f.write(f"\n# Chapter {chapter.chapter_number}: {chapter.title}\n\n")
                        f.write(chapter.content + "\n\n")
                
                print(f"✅ Successfully saved Markdown file: {output_path}")
                
                # Verify the file was created and has content
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"File size: {file_size} bytes")
                    if file_size > 0:
                        print("✅ Markdown file saved successfully with content")
                    else:
                        print("⚠️ Warning: Markdown file saved but appears to be empty")
                else:
                    print("❌ Error: Markdown file was not created")
                    
            except Exception as e:
                print(f"❌ Error saving Markdown file: {str(e)}")
                state["error"] = f"Failed to save Markdown file: {str(e)}"
                return state

        elif format_type == FormatType.DOC:
            print("Saving as DOCX")
            output_filename = f"{filename_base}.docx"
            doc = Document()

            # Title Page
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.font.size = Pt(24)
            title_run.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

            # Table of Contents heading
            toc_heading = doc.add_heading("Table of Contents", level=1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Style the heading to look like a classic book
            run = toc_heading.runs[0]
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Pure black

            # Add some space after heading
            doc.add_paragraph().add_run().add_break()  # Adds a blank line for visual spacing

            # Loop through chapters to add TOC entries
            for chapter in ebook_state.chapters:
                toc_entry = doc.add_paragraph(f"Chapter {chapter.chapter_number}: {chapter.title}")
                toc_entry.style.font.name = 'Times New Roman'
                toc_entry.paragraph_format.left_indent = Inches(0.5)
                toc_entry.paragraph_format.space_after = Pt(6)
                toc_entry.paragraph_format.line_spacing = 1.15

            doc.add_page_break()

            # Chapter content
            for chapter in ebook_state.chapters:
                # Chapter Title

                chapter_num = doc.add_paragraph(f"CHAPTER {chapter.chapter_number}")
                chapter_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = chapter_num.runs[0]
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

                # Chapter title on a new line
                chapter_title = doc.add_paragraph(chapter.title)
                chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = chapter_title.runs[0]
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(14)
                run.font.bold = False
                run.font.color.rgb = RGBColor(0, 0, 0)

                # Add some vertical space
                doc.add_paragraph().add_run().add_break()

                content = chapter.content if hasattr(chapter, 'content') else ''
                paragraphs = content.split("\n\n")

                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue

                    if para.startswith("###"):
                        # Sub-subheading (level 3)
                        subheading = doc.add_heading(para.replace("###", "").strip(), level=3)
                        run = subheading.runs[0]
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    elif para.startswith("##"):
                        # Subheading (level 2)
                        subheading = doc.add_heading(para.replace("##", "").strip(), level=2)
                        run = subheading.runs[0]
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(14)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    else:
                        # Main paragraph text
                        paragraph = doc.add_paragraph(para)
                        run = paragraph.runs[0]
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        paragraph.paragraph_format.line_spacing = 1.5
                        paragraph.paragraph_format.first_line_indent = Inches(0.3)
                        paragraph.paragraph_format.space_after = Pt(6)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                if chapter != ebook_state.chapters[-1]:
                    doc.add_page_break()

            # Save the docx file
            # Ensure output directory exists
            os.makedirs("./output", exist_ok=True)
            
            # Use absolute path for better compatibility
            output_path = os.path.abspath(f"./output/{output_filename}")
            print(f"Saving DOCX file to: {output_path}")
            
            try:
                doc.save(output_path)
                print(f"✅ Successfully saved DOCX file: {output_path}")
                
                # Verify the file was created and has content
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"File size: {file_size} bytes")
                    if file_size > 0:
                        print("✅ File saved successfully with content")
                    else:
                        print("⚠️ Warning: File saved but appears to be empty")
                else:
                    print("❌ Error: File was not created")
                    
            except Exception as e:
                print(f"❌ Error saving DOCX file: {str(e)}")
                print(f"Attempted path: {output_path}")
                # Fallback: try saving in current directory
                fallback_path = f"{output_filename}"
                try:
                    doc.save(fallback_path)
                    print(f"✅ Saved to fallback location: {os.path.abspath(fallback_path)}")
                    output_path = os.path.abspath(fallback_path)
                except Exception as fallback_error:
                    print(f"❌ Fallback save also failed: {str(fallback_error)}")
                    state["error"] = f"Failed to save DOCX file: {str(e)}"
                    return state

        elif format_type == FormatType.PDF:
            print("Saving as PDF")
            output_filename = f"{filename_base}.pdf"

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            pdf.set_font("Times", 'B', 24)
            pdf.multi_cell(0, 12, clean_text(title), align='C')
            pdf.ln(20)

            pdf.set_font("Times", 'B', 18)
            pdf.cell(0, 10, "Table of Contents", ln=True, align='C')
            pdf.ln(10)

            pdf.set_font("Times", '', 12)
            for chapter in ebook_state.chapters:
                toc_line = f"Chapter {chapter.chapter_number}: {chapter.title}"
                pdf.cell(0, 10, clean_text(toc_line), ln=True)
            pdf.add_page()

            for chapter in ebook_state.chapters:
                pdf.set_font("Times", 'B', 16)
                pdf.cell(0, 10, f"CHAPTER {chapter.chapter_number}", ln=True, align='C')

                pdf.set_font("Times", 'B', 14)
                pdf.cell(0, 10, clean_text(chapter.title), ln=True, align='C')
                pdf.ln(5)

                content = getattr(chapter, "content", "")
                paragraphs = content.split("\n\n")

                pdf.set_font("Times", '', 12)
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    cleaned = clean_text(para)
                    if cleaned.startswith("###") or cleaned.startswith("##"):
                        header_text = cleaned.replace("#", "").strip()
                        pdf.set_font("Times", 'B', 12)
                        pdf.multi_cell(0, 10, header_text)
                        pdf.set_font("Times", '', 12)
                        pdf.ln(3)
                    else:
                        pdf.multi_cell(0, 10, cleaned)
                        pdf.ln(5)

                if chapter != ebook_state.chapters[-1]:
                    pdf.add_page()

            # Ensure output directory exists
            os.makedirs("./output", exist_ok=True)
            
            # Use absolute path for better compatibility
            output_path = os.path.abspath(f"./output/{output_filename}")
            print(f"Saving PDF file to: {output_path}")
            
            try:
                pdf.output(output_path)
                print(f"✅ Successfully saved PDF file: {output_path}")
                
                # Verify the file was created and has content
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"File size: {file_size} bytes")
                    if file_size > 0:
                        print("✅ PDF file saved successfully with content")
                    else:
                        print("⚠️ Warning: PDF file saved but appears to be empty")
                else:
                    print("❌ Error: PDF file was not created")
                    
            except Exception as e:
                print(f"❌ Error saving PDF file: {str(e)}")
                print(f"Attempted path: {output_path}")
                # Fallback: try saving in current directory
                fallback_path = f"{output_filename}"
                try:
                    pdf.output(fallback_path)
                    print(f"✅ PDF saved to fallback location: {os.path.abspath(fallback_path)}")
                    output_path = os.path.abspath(fallback_path)
                except Exception as fallback_error:
                    print(f"❌ PDF fallback save also failed: {str(fallback_error)}")
                    state["error"] = f"Failed to save PDF file: {str(e)}"
                    return state


        else:
            print(f"Unsupported format: {format_type}")
            output_filename = f"{filename_base}.md"
            
            # Ensure output directory exists
            os.makedirs("./output", exist_ok=True)
            
            # Use absolute path for better compatibility
            output_path = os.path.abspath(f"./output/{output_filename}")
            print(f"Saving fallback Markdown file to: {output_path}")
            
            try:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(compiled_content)
                print(f"✅ Successfully saved fallback file: {output_path}")
            except Exception as e:
                print(f"❌ Error saving fallback file: {str(e)}")
                state["error"] = f"Failed to save fallback file: {str(e)}"
                return state

        state["output_filename"] = output_filename
        state["ebook_state"] = ebook_state.model_dump()
        print("I am in Format Conversion node")
        return state


class ExportNode:
    """Node to export the e-book to a file"""

    def __init__(self):
        pass

    def __call__(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Export the e-book to a file"""
        output_filename = state.get("output_filename", "ebook.md")
        compiled_content = state.get("compiled_content", "")

        # In a real implementation, we would save the file here
        # For now, we'll just set a flag
        state["export_complete"] = True
        state["file_path"] = f"./output/{output_filename}"
        print("I am in export node")
        return state